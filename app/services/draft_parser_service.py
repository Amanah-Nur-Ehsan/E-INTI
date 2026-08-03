"""Draft parsing and sentence segmentation.

The invariant everything downstream depends on: every char_start/char_end
indexes into `drafts.raw_text`, and `raw_text[start:end] == text` exactly.
To keep that true, each paragraph is normalized *before* its offset is
assigned, raw_text is assembled once, and spaCy runs per block so sentence
offsets compose additively onto the block offset.
"""

import io
import re
import uuid
from dataclasses import asdict, dataclass, field
from datetime import UTC, datetime
from functools import lru_cache
from pathlib import Path

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.logging import get_logger
from app.db.models import Draft
from app.db.models.enums import ParseStatus

log = get_logger(__name__)

BLOCK_SEPARATOR = "\n\n"

#: Tokens after which a period does not end a sentence in academic prose.
ABBREVIATIONS = (
    "et al.",
    "e.g.",
    "i.e.",
    "cf.",
    "vs.",
    "viz.",
    "resp.",
    "approx.",
    "ca.",
    "Fig.",
    "Figs.",
    "Eq.",
    "Eqs.",
    "Ref.",
    "Refs.",
    "Tab.",
    "No.",
    "Sec.",
    "Ch.",
    "pp.",
    "Dr.",
    "Prof.",
    "Mr.",
    "Ms.",
    "St.",
)

_ABBREV_PATTERN = re.compile(
    r"(?:" + "|".join(re.escape(a) for a in ABBREVIATIONS) + r")$", re.IGNORECASE
)
#: A single capital letter followed by a period, i.e. an initial: "J. Smith".
_INITIAL_PATTERN = re.compile(r"\b[A-Z]\.$")
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


@dataclass
class Block:
    """One paragraph of the draft, with its span in raw_text."""

    text: str
    char_start: int
    char_end: int
    paragraph_index: int
    section_title: str | None = None
    is_heading: bool = False


@dataclass
class Sentence:
    text: str
    char_start: int
    char_end: int
    paragraph_index: int
    sentence_index: int
    section_title: str | None = None


@dataclass
class ParsedDraft:
    raw_text: str
    blocks: list[Block] = field(default_factory=list)
    sentences: list[Sentence] = field(default_factory=list)

    def to_json(self) -> dict:
        return {
            "blocks": [asdict(b) for b in self.blocks],
            "sentences": [asdict(s) for s in self.sentences],
        }


def normalize_paragraph(text: str) -> str:
    """Applied before offsets are assigned — never after."""
    text = text.replace(" ", " ").replace("\r", "")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def normalize_with_map(raw: str) -> tuple[str, list[int]]:
    """Like `normalize_paragraph`, but also returns the source index each
    output character came from: `src[i]` is the position in `raw` that
    produced `normalized[i]`.

    This is what lets the export stage invert the normalization: a claim
    offset lands in `raw_text` (normalized space), and this map is how it
    finds its way back to a character position in the original docx
    paragraph text (`raw` space) without ever persisting a mapping.

    Must stay in exact lockstep with `normalize_paragraph` -- the two are
    unit-tested against each other over every fixture paragraph.
    """
    out: list[str] = []
    src: list[int] = []
    pending_space_src: int | None = None

    for i, ch in enumerate(raw):
        if ch == "\r":
            continue
        if ch == "\xa0":
            ch = " "
        if ch in (" ", "\t"):
            if pending_space_src is None:
                pending_space_src = i
            continue
        if pending_space_src is not None:
            out.append(" ")
            src.append(pending_space_src)
            pending_space_src = None
        out.append(ch)
        src.append(i)

    if pending_space_src is not None:
        out.append(" ")
        src.append(pending_space_src)

    # Python's str.strip() strips any whitespace, not just the space/tab
    # runs collapsed above -- a bare newline at either edge (from a \r\n
    # that lost its \r) must go too, matching normalize_paragraph exactly.
    start = 0
    while start < len(out) and out[start].isspace():
        start += 1
    end = len(out)
    while end > start and out[end - 1].isspace():
        end -= 1

    return "".join(out[start:end]), src[start:end]


def _assemble(raw_blocks: list[tuple[str, bool, str | None]]) -> ParsedDraft:
    """Build raw_text and block offsets in one pass.

    raw_blocks items are (text, is_heading, explicit_section_title).
    """
    pieces: list[str] = []
    blocks: list[Block] = []
    cursor = 0
    current_section: str | None = None

    for index, (text, is_heading, explicit_section) in enumerate(raw_blocks):
        normalized = normalize_paragraph(text)
        if not normalized:
            continue

        if cursor > 0:
            pieces.append(BLOCK_SEPARATOR)
            cursor += len(BLOCK_SEPARATOR)

        start = cursor
        pieces.append(normalized)
        cursor += len(normalized)

        if is_heading:
            current_section = normalized

        blocks.append(
            Block(
                text=normalized,
                char_start=start,
                char_end=cursor,
                paragraph_index=index,
                section_title=explicit_section or (normalized if is_heading else current_section),
                is_heading=is_heading,
            )
        )

    return ParsedDraft(raw_text="".join(pieces), blocks=blocks)


def iter_docx_paragraphs(document) -> list:
    """The exact paragraph walk that produced `raw_blocks` in `parse_docx`:
    body paragraphs first, then every table cell's paragraphs, in document
    order and depth-first per table.

    Export replays this same walk to turn a block's `paragraph_index` back
    into a python-docx `Paragraph` object. The two call sites living in one
    function, in one module, is what keeps them from drifting apart — if
    this walk ever changes, both the parser and the exporter change with it.

    Two consequences worth knowing rather than "fixing": a horizontally
    merged cell yields its paragraphs once per grid column it spans (both
    sides duplicate identically, so indices still agree), and paragraphs
    inside a nested table are not visited by either side.
    """
    from docx.table import Table

    flat = list(document.paragraphs)
    for table in document.tables:
        if not isinstance(table, Table):
            continue
        for row in table.rows:
            for cell in row.cells:
                flat.extend(cell.paragraphs)
    return flat


def _para_text(paragraph) -> str:
    """The paragraph's text, exactly as `paragraph.text` would report it —
    except that `paragraph.text` silently omits any run living inside a
    `<w:ins>` (tracked-insert) element. A previously-exported document has
    such runs, so validation must never read `.text` directly; this walks
    `.//w:r` instead, which includes them.

    `<w:del>` (tracked-delete) content is excluded: a deleted run's text
    is not part of the paragraph as Word will show it once changes are
    accepted, so it must not count towards offsets either.
    """
    from docx.oxml.ns import qn

    parts = []
    for run_el in paragraph._p.xpath(".//w:r[not(ancestor::w:del)]"):
        for t in run_el.findall(qn("w:t")):
            parts.append(t.text or "")
        for _ in run_el.findall(qn("w:tab")):
            parts.append("\t")
        for _ in run_el.findall(qn("w:br")) + run_el.findall(qn("w:cr")):
            parts.append("\n")
    return "".join(parts)


#: Smallest valid PNG (1x1 transparent). Substituted for embedded images
#: whose bytes fail their CRC check, so the package's relationship graph
#: stays intact while the unreadable pixels are dropped.
_PLACEHOLDER_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4"
    "890000000a49444154789c63000100000500010d0a2db40000000049454e44ae426082"
)


def _repair_docx_zip(path: Path) -> io.BytesIO:
    """Rewrite a .docx whose zip has corrupt members, dropping their bytes.

    Word files that have been through email, cloud sync, or a partial
    upload routinely carry a truncated image whose CRC no longer matches,
    and python-docx reads *every* part eagerly -- so one bad PNG makes an
    otherwise perfectly readable document unparseable. We only ever want
    the text, so a corrupt part is replaced (images with a 1x1 placeholder,
    anything else with empty bytes) rather than dropped outright, which
    would break the relationships that reference it.
    """
    import zipfile

    repaired = io.BytesIO()
    damaged: list[str] = []

    with zipfile.ZipFile(path) as source:
        with zipfile.ZipFile(repaired, "w", zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                try:
                    data = source.read(info.filename)
                except (zipfile.BadZipFile, EOFError, OSError):
                    damaged.append(info.filename)
                    lowered = info.filename.lower()
                    is_image = lowered.startswith("word/media/") or lowered.endswith(
                        (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".emf", ".wmf")
                    )
                    data = _PLACEHOLDER_PNG if is_image else b""
                target.writestr(info.filename, data)

    if not damaged:
        # The zip failed for a reason we haven't handled; let the caller's
        # retry surface the original error rather than pretending we fixed it.
        raise ValueError("docx zip reported corrupt but every member read cleanly")

    log.warning("docx_corrupt_parts_replaced", path=str(path), parts=damaged)
    repaired.seek(0)
    return repaired


def parse_docx(path: Path) -> ParsedDraft:
    import zipfile

    from docx import Document

    try:
        document = Document(str(path))
    except zipfile.BadZipFile:
        # Text is all we need; a damaged embedded image must not cost the
        # user their whole upload.
        document = Document(_repair_docx_zip(path))

    raw_blocks: list[tuple[str, bool, str | None]] = []

    for paragraph in iter_docx_paragraphs(document):
        style = (paragraph.style.name or "") if paragraph.style is not None else ""
        is_heading = style.startswith("Heading") or style == "Title"
        raw_blocks.append((paragraph.text, is_heading, None))

    return _assemble(raw_blocks)


def parse_markdown(path: Path) -> ParsedDraft:
    raw_blocks: list[tuple[str, bool, str | None]] = []
    for chunk in re.split(r"\n\s*\n", path.read_text(encoding="utf-8")):
        chunk = chunk.strip()
        if not chunk:
            continue
        heading = _MD_HEADING.match(chunk)
        if heading:
            raw_blocks.append((heading.group(2), True, None))
        else:
            raw_blocks.append((chunk.replace("\n", " "), False, None))
    return _assemble(raw_blocks)


def parse_plaintext(path: Path) -> ParsedDraft:
    raw_blocks = [
        (chunk.replace("\n", " "), False, None)
        for chunk in re.split(r"\n\s*\n", path.read_text(encoding="utf-8"))
        if chunk.strip()
    ]
    return _assemble(raw_blocks)


def parse_file(path: Path) -> ParsedDraft:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        parsed = parse_docx(path)
    elif suffix in (".md", ".markdown"):
        parsed = parse_markdown(path)
    elif suffix in (".txt", ".text"):
        parsed = parse_plaintext(path)
    else:
        raise ValueError(f"Unsupported draft format: {path.suffix}")

    parsed.sentences = segment(parsed)
    return parsed


@lru_cache
def get_nlp():
    import spacy
    from spacy.language import Language

    @Language.component("academic_sentence_fixer")
    def academic_sentence_fixer(doc):
        for index, token in enumerate(doc[:-1]):
            prefix = doc[: index + 1].text
            following = doc[index + 1]
            if _ABBREV_PATTERN.search(prefix) or _INITIAL_PATTERN.search(prefix):
                following.is_sent_start = False
            # "3.5 percent" — a period between digits is a decimal point.
            if token.text == "." and index > 0:
                if doc[index - 1].text.isdigit() and following.text[:1].isdigit():
                    following.is_sent_start = False
        return doc

    nlp = spacy.load(get_settings().spacy_model, exclude=["ner", "lemmatizer"])
    if "academic_sentence_fixer" not in nlp.pipe_names:
        nlp.add_pipe("academic_sentence_fixer", before="parser")
    return nlp


def segment(parsed: ParsedDraft) -> list[Sentence]:
    """Segment each block separately so offsets compose additively."""
    nlp = get_nlp()
    sentences: list[Sentence] = []

    body_blocks = [b for b in parsed.blocks if not b.is_heading]
    docs = nlp.pipe([b.text for b in body_blocks])

    for block, doc in zip(body_blocks, docs, strict=True):
        for sentence_index, sent in enumerate(doc.sents):
            text = sent.text.strip()
            if not text:
                continue
            # Re-anchor onto the block text to absorb spaCy's own whitespace trimming.
            offset = block.text.find(text, sent.start_char if sent.start_char < len(block.text) else 0)
            if offset < 0:
                offset = sent.start_char
            sentences.append(
                Sentence(
                    text=text,
                    char_start=block.char_start + offset,
                    char_end=block.char_start + offset + len(text),
                    paragraph_index=block.paragraph_index,
                    sentence_index=sentence_index,
                    section_title=block.section_title,
                )
            )

    return sentences


def local_context(sentences: list[Sentence], index: int) -> str:
    """Previous + current + next sentence — what retrieval actually queries on."""
    window = sentences[max(0, index - 1) : index + 2]
    return " ".join(s.text for s in window)


def parse_and_store_draft(session: Session, draft_id: uuid.UUID) -> dict:
    """Stage body: parse the uploaded file and persist text, blocks, sentences."""
    draft = session.get(Draft, draft_id)
    if draft is None:
        raise ValueError(f"Draft {draft_id} not found")

    path = Path(draft.storage_path)
    if not path.exists():
        raise FileNotFoundError(f"Draft file missing: {path}")

    try:
        parsed = parse_file(path)
    except Exception as exc:
        draft.parse_status = ParseStatus.FAILED
        draft.parse_error = f"{type(exc).__name__}: {exc}"
        session.commit()
        raise

    draft.raw_text = parsed.raw_text
    draft.parsed_content = parsed.to_json()
    draft.parse_status = ParseStatus.PARSED
    draft.parse_error = None
    draft.parsed_at = datetime.now(UTC)
    session.commit()

    return {"blocks": len(parsed.blocks), "sentences": len(parsed.sentences)}
