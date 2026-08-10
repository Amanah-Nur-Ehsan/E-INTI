"""Produces the revised DOCX: accepted citations spliced into the source
paragraphs, plus a generated bibliography.

Always opens the pristine file at `draft.storage_path` — never a
previous export's output — because `paragraph.text` (and therefore the
offset-mapping validation in `resolve_insertion_point`) silently omits
text already wrapped in `<w:ins>`, which would corrupt validation against
a document that already carries tracked changes.
"""

import re
import uuid
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO
from typing import Literal

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from app.db.models.enums import InsertionMode
from app.services.citation_formatting_service import join_in_text
from app.services.draft_parser_service import iter_docx_paragraphs
from app.services.export.bundle import ExportBundle
from app.services.export.docx_ops import (
    REVISION_ID_BASE,
    clone_run_formatting,
    insert_node_at,
    resolve_insertion_point,
    wrap_tracked_insert,
)

#: Same pattern as claim_detection_service.py's _INTRO_HEADING -- kept as an
#: independent copy rather than a shared import, since this module reasons
#: about export-time paragraph objects while that one reasons about
#: parsed-content dicts, and the two have no other reason to depend on each
#: other.
_INTRO_HEADING = re.compile(r"^\s*(?:\d+[.)]?\s*|[ivxlcdm]+[.)]\s*)?introduction\b", re.IGNORECASE)
_KEYWORDS_LINE = re.compile(r"^\s*(?:key\s*words?|kata\s*kunci)\s*[:\-]?", re.IGNORECASE)

InsertionOutcomeStatus = Literal[
    "INSERTED", "POSITION_MISMATCH", "NO_BLOCK", "EMPTY_CITATION"
]


@dataclass(frozen=True)
class InsertionOutcome:
    claim_id: uuid.UUID
    status: InsertionOutcomeStatus
    citation_text: str | None
    paragraph_index: int | None


TRACKED_CHANGE_AUTHOR = "CitationRecommender"


def _blocks_by_paragraph_index(bundle: ExportBundle) -> dict[int, object]:
    return {block.paragraph_index: block for block in bundle.blocks}


def _build_citation_node(source_run_el, text: str, mode: str):
    """The element to splice in, varying only by insertion mode.

    Always yellow-highlighted: an inserted citation should stay visibly
    identifiable as an INTI addition even after tracked changes are
    accepted and the `<w:ins>` markup is gone.
    """
    if mode == InsertionMode.PLACEHOLDER:
        return clone_run_formatting(source_run_el, f" [CITATION: {text}]", highlight=True)
    return clone_run_formatting(source_run_el, f" {text}", highlight=True)


def write_docx(
    bundle: ExportBundle,
    mode: str = InsertionMode.TRACKED_CHANGES,
    author: str = TRACKED_CHANGE_AUTHOR,
    when: datetime | None = None,
    include_audit_report: bool = False,
) -> tuple[bytes, list[InsertionOutcome]]:
    """Returns the revised document bytes plus a per-claim outcome list.

    A claim whose paragraph or sentence text no longer matches what was
    stored at analysis time is skipped (POSITION_MISMATCH), never guessed.
    """
    if not bundle.draft.storage_path.lower().endswith(".docx"):
        raise ValueError(
            "DOCX export requires a .docx source draft; use the markdown export "
            "for .md/.txt drafts instead"
        )

    document = Document(bundle.draft.storage_path)
    flat_paragraphs = iter_docx_paragraphs(document)
    blocks_by_index = _blocks_by_paragraph_index(bundle)
    when = when or datetime.now(UTC)

    outcomes: list[InsertionOutcome] = []

    # Group by paragraph, then process each paragraph's insertions in
    # descending char_start order: an insertion shifts every later offset
    # in the same paragraph, so ascending order would corrupt everything
    # after the first one. Different paragraphs are independent and can be
    # processed in any order relative to each other.
    by_paragraph: dict[int, list] = {}
    for item in bundle.accepted:
        block = blocks_by_index.get(item.claim.paragraph_index)
        if block is None:
            outcomes.append(
                InsertionOutcome(
                    claim_id=item.claim.id,
                    status="NO_BLOCK",
                    citation_text=None,
                    paragraph_index=item.claim.paragraph_index,
                )
            )
            continue
        by_paragraph.setdefault(item.claim.paragraph_index, []).append(item)

    revision_counter = 0
    for _paragraph_index, items in by_paragraph.items():
        items.sort(key=lambda i: i.claim.char_start or 0, reverse=True)

        # Resolve every insertion point against the *pristine* paragraph
        # before mutating anything. Resolving one at a time, interleaved
        # with insertion, would fail: the first insertion changes what
        # _para_text(paragraph) returns (it walks .//w:r, which includes
        # our own freshly-added <w:ins> runs), so validating the second
        # item against the now-mutated paragraph would see extra text it
        # doesn't recognize and report a false POSITION_MISMATCH. The
        # (run_element, offset_in_run) pairs captured here stay valid
        # through the later mutations: descending order means an earlier
        # (lower-offset) item's target run is only ever truncated *after*
        # its own offset, never before it, even when two items share a run.
        resolved: list[tuple] = []
        for item in items:
            block = blocks_by_index[item.claim.paragraph_index]
            citations = [bundle.context.in_text(ref.id) for ref in item.references]
            joined = join_in_text(citations)
            if not joined:
                outcomes.append(
                    InsertionOutcome(
                        claim_id=item.claim.id,
                        status="EMPTY_CITATION",
                        citation_text=None,
                        paragraph_index=item.claim.paragraph_index,
                    )
                )
                resolved.append((item, None, joined))
                continue

            point = resolve_insertion_point(
                document,
                flat_paragraphs,
                block,
                item.claim.char_start,
                item.claim.char_end,
                item.claim.sentence_text,
            )
            resolved.append((item, point, joined))

        for item, point, joined in resolved:
            if point is None:
                if joined:  # EMPTY_CITATION already recorded above
                    outcomes.append(
                        InsertionOutcome(
                            claim_id=item.claim.id,
                            status="POSITION_MISMATCH",
                            citation_text=joined,
                            paragraph_index=item.claim.paragraph_index,
                        )
                    )
                continue

            node = _build_citation_node(point.run_element, joined, mode)
            if mode == InsertionMode.TRACKED_CHANGES:
                revision_counter += 1
                node = wrap_tracked_insert(
                    node, author=author, when=when, revision_id=revision_counter
                )
            insert_node_at(point, node)

            outcomes.append(
                InsertionOutcome(
                    claim_id=item.claim.id,
                    status="INSERTED",
                    citation_text=joined,
                    paragraph_index=item.claim.paragraph_index,
                )
            )

    revision_counter = _append_bibliography(
        document, bundle, mode=mode, author=author, when=when, revision_counter=revision_counter
    )
    revision_counter = _append_sdg_keyword(
        document,
        bundle,
        flat_paragraphs,
        mode=mode,
        author=author,
        when=when,
        revision_counter=revision_counter,
    )
    if include_audit_report:
        _append_audit_table(document, bundle)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue(), outcomes


def _append_audit_table(document, bundle: ExportBundle) -> None:
    """A plain (non-tracked) table summarising every claim that needed a
    citation -- section, category, existing-citation status, the chosen
    reference's score and verdict, and the user's decision. This is a
    report for the researcher, not manuscript content, so it is never
    wrapped in <w:ins> regardless of insertion mode.
    """
    from app.services.export.audit import build_audit_rows

    rows = build_audit_rows(bundle)
    if not rows:
        return

    document.add_page_break()
    document.add_heading("Citation Audit Report", level=1)

    columns = [
        ("Section", lambda r: r.section_title or ""),
        ("Claim", lambda r: (r.claim_text or "")[:120]),
        ("Type", lambda r: r.claim_type or ""),
        ("Existing citation", lambda r: r.existing_citation_status),
        ("Score", lambda r: f"{r.score_percentage:.1f}%" if r.score_percentage is not None else ""),
        ("Verdict", lambda r: r.verdict or ""),
        ("Decision", lambda r: r.user_decision),
        ("Status", lambda r: r.insertion_status),
    ]

    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"  # a built-in style present in every default template
    for cell, (header, _getter) in zip(table.rows[0].cells, columns, strict=True):
        cell.text = header

    for row in rows:
        cells = table.add_row().cells
        for cell, (_header, getter) in zip(cells, columns, strict=True):
            cell.text = str(getter(row))


def _append_bibliography(
    document,
    bundle: ExportBundle,
    *,
    mode: str,
    author: str,
    when: datetime,
    revision_counter: int,
) -> int:
    """Appends the References section; returns the updated revision counter
    so ids stay unique and monotonic across the whole document, citations
    and bibliography alike.
    """
    entries = bundle.context.bibliography()
    if not entries:
        return revision_counter

    document.add_page_break()
    heading = document.add_heading("References", level=1)
    if mode == InsertionMode.TRACKED_CHANGES:
        revision_counter += 1
        _wrap_paragraph_runs_as_tracked(
            heading, author=author, when=when, revision_id=revision_counter
        )

    for _reference_id, segments in entries:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        for segment in segments:
            run = paragraph.add_run(segment.text)
            run.italic = segment.italic
            # Highlighted like every other INTI addition -- see
            # _build_citation_node's docstring for why yellow specifically.
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        if mode == InsertionMode.TRACKED_CHANGES:
            revision_counter += 1
            _wrap_paragraph_runs_as_tracked(
                paragraph, author=author, when=when, revision_id=revision_counter
            )

    return revision_counter


def _append_sdg_keyword(
    document,
    bundle: ExportBundle,
    flat_paragraphs: list,
    *,
    mode: str,
    author: str,
    when: datetime,
    revision_counter: int,
) -> int:
    """Adds the classified SDG keyword to the paper's Keywords line,
    highlighted like every other INTI addition.

    A no-op when the classifier declined (`sdg_number` null -- see
    `sdg_classification_service.classify_draft`'s `fits` field) or somehow
    has a goal with no matched keyword: nothing genuine to add, so nothing
    is written.
    """
    keyword = bundle.draft.sdg_keyword
    if not keyword or not bundle.draft.sdg_number:
        return revision_counter

    keywords_block = next((b for b in bundle.blocks if _KEYWORDS_LINE.match(b.text)), None)

    if keywords_block is not None and keywords_block.paragraph_index < len(flat_paragraphs):
        paragraph = flat_paragraphs[keywords_block.paragraph_index]
        run_elements = paragraph._p.xpath(".//w:r[not(ancestor::w:del)]")
        source_el = run_elements[-1] if run_elements else None
        new_run = (
            clone_run_formatting(source_el, f"; {keyword}", highlight=True)
            if source_el is not None
            else clone_run_formatting(OxmlElement("w:r"), f"; {keyword}", highlight=True)
        )
        node = new_run
        if mode == InsertionMode.TRACKED_CHANGES:
            revision_counter += 1
            node = wrap_tracked_insert(new_run, author=author, when=when, revision_id=revision_counter)
        paragraph._p.append(node)
        return revision_counter

    # No existing Keywords line: add one right before the Introduction
    # heading (the same cut point claim detection uses to skip the
    # abstract), so it lands with the rest of the front matter rather than
    # mid-body. Falls back to the very top of the document if there's no
    # Introduction heading either.
    intro_block = next(
        (b for b in bundle.blocks if b.is_heading and _INTRO_HEADING.match(b.text)), None
    )
    if not flat_paragraphs:
        return revision_counter
    anchor_index = intro_block.paragraph_index if intro_block is not None else 0
    if anchor_index >= len(flat_paragraphs):
        return revision_counter
    anchor_paragraph = flat_paragraphs[anchor_index]

    new_paragraph = anchor_paragraph.insert_paragraph_before()
    run = new_paragraph.add_run(f"Keywords: {keyword}")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if mode == InsertionMode.TRACKED_CHANGES:
        revision_counter += 1
        _wrap_paragraph_runs_as_tracked(
            new_paragraph, author=author, when=when, revision_id=revision_counter
        )
    return revision_counter


def _wrap_paragraph_runs_as_tracked(paragraph, *, author: str, when: datetime, revision_id: int) -> None:
    """Wrap every run already in a freshly-added paragraph in a single
    `<w:ins>`, so a generated heading or bibliography entry shows as one
    reviewable insertion rather than plain text.
    """
    p_el = paragraph._p
    run_elements = p_el.xpath("./w:r")
    if not run_elements:
        return
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(REVISION_ID_BASE + revision_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), when.strftime("%Y-%m-%dT%H:%M:%SZ"))
    first = run_elements[0]
    first.addprevious(ins)
    for run_el in run_elements:
        ins.append(run_el)
