"""Unit tests for the oxml run-splitting primitives, using small synthetic
documents built in-memory so each behaviour is isolated.
"""

import io

from docx import Document
from docx.oxml.ns import qn

from app.services.draft_parser_service import Block, iter_docx_paragraphs
from app.services.export.docx_ops import (
    InsertionPoint,
    _para_text,
    _run_text,
    clone_run_formatting,
    find_insertion_offset,
    insert_node_at,
    resolve_insertion_point,
    wrap_tracked_insert,
)


def _build_doc(paragraphs: list[str]) -> Document:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


def _reload(doc: Document) -> Document:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


def test_find_insertion_offset_lands_before_terminal_period():
    normalized = "A claim ending here. Another sentence follows."
    offset = find_insertion_offset(normalized, 0, 20)  # "A claim ending here."
    assert normalized[:offset] == "A claim ending here"


def test_find_insertion_offset_handles_bracket_citation():
    normalized = "Something was shown [1]."
    offset = find_insertion_offset(normalized, 0, len(normalized))
    assert normalized[:offset] == "Something was shown [1]"


def test_find_insertion_offset_no_terminal_punctuation():
    normalized = "A sentence with no ending"
    offset = find_insertion_offset(normalized, 0, len(normalized))
    assert offset == len(normalized)


def test_insert_node_splits_run_preserving_bold_on_both_halves():
    doc = _build_doc(["placeholder"])
    paragraph = doc.paragraphs[0]
    paragraph.runs[0].text = "Hello world today"
    paragraph.runs[0].bold = True
    run_el = paragraph._p.xpath(".//w:r")[0]

    point = InsertionPoint(paragraph=paragraph, run_element=run_el, offset_in_run=5)
    node = clone_run_formatting(run_el, " INSERTED")
    insert_node_at(point, node)

    reloaded = _reload(doc)
    runs = reloaded.paragraphs[0]._p.xpath(".//w:r")
    assert len(runs) == 3
    assert "".join(_run_text(r) for r in runs) == "Hello INSERTED world today"
    for r in runs:
        rpr = r.find(qn("w:rPr"))
        assert rpr is not None
        assert rpr.find(qn("w:b")) is not None


def test_insert_node_at_start_of_run_uses_addprevious():
    doc = _build_doc(["Hello world"])
    paragraph = doc.paragraphs[0]
    run_el = paragraph._p.xpath(".//w:r")[0]
    point = InsertionPoint(paragraph=paragraph, run_element=run_el, offset_in_run=0)
    node = clone_run_formatting(run_el, "PREFIX ")
    insert_node_at(point, node)

    reloaded = _reload(doc)
    runs = reloaded.paragraphs[0]._p.xpath(".//w:r")
    assert "".join(_run_text(r) for r in runs) == "PREFIX Hello world"


def test_insert_node_at_end_of_run_uses_addnext():
    doc = _build_doc(["Hello world"])
    paragraph = doc.paragraphs[0]
    run_el = paragraph._p.xpath(".//w:r")[0]
    point = InsertionPoint(
        paragraph=paragraph, run_element=run_el, offset_in_run=len(_run_text(run_el))
    )
    node = clone_run_formatting(run_el, " SUFFIX")
    insert_node_at(point, node)

    reloaded = _reload(doc)
    runs = reloaded.paragraphs[0]._p.xpath(".//w:r")
    assert "".join(_run_text(r) for r in runs) == "Hello world SUFFIX"


def test_wrap_tracked_insert_produces_valid_ins_element():
    doc = _build_doc(["Hello world"])
    paragraph = doc.paragraphs[0]
    run_el = paragraph._p.xpath(".//w:r")[0]  # stays attached; only the NEW node is wrapped
    node = clone_run_formatting(run_el, " (Smith, 2023)")
    ins = wrap_tracked_insert(run_el=node, author="CitationRecommender", when=None, revision_id=1)

    insert_point = InsertionPoint(paragraph=paragraph, run_element=run_el, offset_in_run=11)
    insert_node_at(insert_point, ins)

    reloaded = _reload(doc)
    ins_elements = reloaded.paragraphs[0]._p.xpath(".//w:ins")
    assert len(ins_elements) == 1
    assert ins_elements[0].get(qn("w:author")) == "CitationRecommender"
    assert ins_elements[0].get(qn("w:id")) is not None
    assert "(Smith, 2023)" in _para_text(reloaded.paragraphs[0])


def test_para_text_excludes_deleted_runs():
    doc = _build_doc(["Hello world"])
    paragraph = doc.paragraphs[0]
    run_el = paragraph._p.xpath(".//w:r")[0]

    from docx.oxml import OxmlElement

    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), "1")
    del_el.set(qn("w:author"), "someone")
    deleted_run = clone_run_formatting(run_el, "DELETED")
    # <w:delText> is what a real deleted run uses, but for this test we only
    # need `_para_text` to skip runs under w:del entirely, regardless of tag.
    del_el.append(deleted_run)
    paragraph._p.append(del_el)

    reloaded = _reload(doc)
    assert "DELETED" not in _para_text(reloaded.paragraphs[0])
    assert _para_text(reloaded.paragraphs[0]) == "Hello world"


def test_resolve_insertion_point_detects_paragraph_mismatch():
    doc = _build_doc(["Original sentence text here for testing."])
    reloaded = _reload(doc)
    flat = iter_docx_paragraphs(reloaded)

    block = Block(
        text="A completely different paragraph.",
        char_start=0,
        char_end=34,
        paragraph_index=0,
    )
    point = resolve_insertion_point(
        reloaded, flat, block, 0, 34, "A completely different paragraph."
    )
    assert point is None


def test_resolve_insertion_point_detects_sentence_mismatch_within_matching_paragraph():
    text = "First sentence here. Second sentence follows."
    doc = _build_doc([text])
    reloaded = _reload(doc)
    flat = iter_docx_paragraphs(reloaded)

    block = Block(text=text, char_start=0, char_end=len(text), paragraph_index=0)
    # Claim to a sentence span that does not match what's actually there.
    point = resolve_insertion_point(reloaded, flat, block, 0, 20, "A sentence that never existed")
    assert point is None


def test_resolve_insertion_point_succeeds_on_matching_paragraph():
    text = "First sentence here. Second sentence follows."
    doc = _build_doc([text])
    reloaded = _reload(doc)
    flat = iter_docx_paragraphs(reloaded)

    block = Block(text=text, char_start=0, char_end=len(text), paragraph_index=0)
    point = resolve_insertion_point(reloaded, flat, block, 0, 20, "First sentence here.")
    assert point is not None
    assert point.paragraph is flat[0]
