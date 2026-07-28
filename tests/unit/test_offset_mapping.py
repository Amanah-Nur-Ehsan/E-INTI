"""The load-bearing suite for phase 6: normalize_with_map must agree with
normalize_paragraph exactly, and iter_docx_paragraphs must reproduce the
paragraph_index identity parse_docx relied on when it built raw_text.
Everything the export stage does rests on both holding for every
paragraph of every fixture, and of the real validation paper.
"""

from pathlib import Path

import pytest
from docx import Document

from app.services.draft_parser_service import (
    iter_docx_paragraphs,
    normalize_paragraph,
    normalize_with_map,
    parse_file,
)

FIXTURES = Path(__file__).parent.parent / "fixtures"
REAL_PAPER = (
    Path(__file__).parent.parent.parent
    / "[Draft] Listening Like a Clinician- Concept-Aligned Explainable AI for "
    "Trustworthy Classification of Respiratory and Cardiac Sounds.docx"
)


@pytest.mark.parametrize(
    "raw",
    [
        "",
        "   ",
        "plain text",
        "  leading and trailing  ",
        "double  internal   spaces",
        "tabs\tand\tmore\ttabs",
        "\r\ncarriage\rreturns\r\n",
        "nbsp\xa0here\xa0and\xa0here",
        "mixed \xa0 \t  whitespace   runs",
        "et al. (2024) showed X.",
        "single",
        " ",
        "\xa0",
        "a\xa0\xa0\xa0b",
    ],
)
def test_normalize_with_map_matches_normalize_paragraph(raw):
    normalized, src = normalize_with_map(raw)
    assert normalized == normalize_paragraph(raw)


@pytest.mark.parametrize(
    "raw",
    [
        "plain text here",
        "  leading trailing  ",
        "double  spaces   collapse",
        "nbsp\xa0separated\xa0words",
        "tabs\tinstead\tof\tspaces",
    ],
)
def test_src_map_is_strictly_increasing_and_in_bounds(raw):
    normalized, src = normalize_with_map(raw)
    assert len(src) == len(normalized)
    for index in src:
        assert 0 <= index < len(raw)
    for a, b in zip(src, src[1:], strict=False):
        assert a < b


@pytest.mark.parametrize(
    "raw",
    [
        "plain text here",
        "double  spaces   collapse to one",
        "nbsp\xa0words\xa0here",
        "trailing punctuation.",
        "et al. (2024) claim ends here.",
    ],
)
def test_src_map_recovers_correct_source_characters(raw):
    """src[i] must point at a source character that actually produced
    normalized[i] -- for non-space output chars this is an exact identity,
    not just "some earlier index"."""
    normalized, src = normalize_with_map(raw)
    for i, out_ch in enumerate(normalized):
        if out_ch != " ":
            assert raw[src[i]] == out_ch


@pytest.mark.parametrize(
    "filename",
    ["sample_draft.docx"],
)
def test_normalize_with_map_matches_on_every_fixture_paragraph(filename):
    document = Document(str(FIXTURES / filename))
    for paragraph in iter_docx_paragraphs(document):
        normalized, _src = normalize_with_map(paragraph.text)
        assert normalized == normalize_paragraph(paragraph.text)


def test_iter_docx_paragraphs_matches_parser_block_assignment_on_fixture():
    parsed = parse_file(FIXTURES / "sample_draft.docx")
    document = Document(str(FIXTURES / "sample_draft.docx"))
    flat = iter_docx_paragraphs(document)

    for block in parsed.blocks:
        assert block.paragraph_index < len(flat)
        assert normalize_paragraph(flat[block.paragraph_index].text) == block.text


@pytest.mark.skipif(not REAL_PAPER.exists(), reason="real validation paper not present locally")
def test_offset_mapping_holds_on_the_real_29k_char_paper():
    parsed = parse_file(REAL_PAPER)
    document = Document(str(REAL_PAPER))
    flat = iter_docx_paragraphs(document)

    assert len(flat) >= max(b.paragraph_index for b in parsed.blocks) + 1

    mismatches = 0
    for block in parsed.blocks:
        if normalize_paragraph(flat[block.paragraph_index].text) != block.text:
            mismatches += 1
    assert mismatches == 0

    for paragraph in flat:
        normalized, _src = normalize_with_map(paragraph.text)
        assert normalized == normalize_paragraph(paragraph.text)

    for sentence in parsed.sentences:
        assert parsed.raw_text[sentence.char_start : sentence.char_end] == sentence.text
