"""POST /drafts/{id}/exports and the download endpoint, across all
four formats, through the real HTTP layer.
"""

import pytest
from docx import Document

pytestmark = pytest.mark.integration


async def _analyzed_draft(client, db_session, seeded_draft):
    draft_id = seeded_draft
    await client.post(f"/api/v1/drafts/{draft_id}/analysis/run")

    claims = (await client.get(f"/api/v1/drafts/{draft_id}/claims?needs_citation=true")).json()
    for claim in claims:
        recs = (await client.get(f"/api/v1/claims/{claim['id']}/recommendations")).json()
        if recs:
            await client.post(f"/api/v1/recommendations/{recs[0]['id']}/accept", json={})
    return draft_id


@pytest.mark.parametrize("fmt", ["docx", "md", "csv", "json"])
async def test_export_and_download_round_trip(client, db_session, seeded_draft, fmt):
    draft_id = await _analyzed_draft(client, db_session, seeded_draft)

    resp = await client.post(f"/api/v1/drafts/{draft_id}/exports", json={"format": fmt})
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["format"] == fmt
    assert body["byte_size"] > 0
    assert body["filename"].endswith(f".{fmt}")

    download = await client.get(f"/api/v1/exports/{body['id']}/download")
    assert download.status_code == 200
    assert len(download.content) == body["byte_size"]


async def test_export_filename_is_prefixed_with_the_source_paper_name(client, db_session, seeded_draft):
    """The download's display name is "[INTI] <paper name>.<ext>", derived
    from the draft's original filename -- not the UUID-based storage name
    on disk (see app/services/export/__init__.py's run_export)."""
    draft_id = await _analyzed_draft(client, db_session, seeded_draft)

    resp = await client.post(f"/api/v1/drafts/{draft_id}/exports", json={"format": "docx"})
    assert resp.status_code == 201, resp.text
    body = resp.json()

    assert body["filename"].startswith("[INTI] ")
    assert body["filename"] == "[INTI] sample_draft.docx"

    download = await client.get(f"/api/v1/exports/{body['id']}/download")
    # Starlette encodes non-ASCII-safe characters (the brackets, the space)
    # into the UTF-8 filename* form per RFC 5987/6266 rather than the plain
    # filename= form -- still the same name, just percent-encoded.
    disposition = download.headers["content-disposition"]
    assert "filename*=utf-8''" in disposition
    assert "%5BINTI%5D" in disposition
    assert "sample_draft.docx" in disposition


async def test_docx_export_reports_inserted_count(client, db_session, seeded_draft):
    draft_id = await _analyzed_draft(client, db_session, seeded_draft)
    resp = await client.post(f"/api/v1/drafts/{draft_id}/exports", json={"format": "docx"})
    body = resp.json()
    assert body["inserted_count"] > 0
    assert body["mismatch_count"] == 0
    assert len(body["outcomes"]) == body["inserted_count"] + body["mismatch_count"]


async def test_docx_export_includes_audit_table_when_requested(
    client, db_session, seeded_draft, tmp_path
):
    draft_id = await _analyzed_draft(client, db_session, seeded_draft)
    resp = await client.post(
        f"/api/v1/drafts/{draft_id}/exports",
        json={"format": "docx", "include_audit_report": True},
    )
    export_id = resp.json()["id"]
    download = await client.get(f"/api/v1/exports/{export_id}/download")

    out = tmp_path / "export.docx"
    out.write_bytes(download.content)
    doc = Document(str(out))
    assert len(doc.tables) == 1
    assert [c.text for c in doc.tables[0].rows[0].cells][0] == "Section"


async def test_docx_export_excludes_audit_table_when_not_requested(
    client, db_session, seeded_draft, tmp_path
):
    draft_id = await _analyzed_draft(client, db_session, seeded_draft)
    resp = await client.post(
        f"/api/v1/drafts/{draft_id}/exports",
        json={"format": "docx", "include_audit_report": False},
    )
    export_id = resp.json()["id"]
    download = await client.get(f"/api/v1/exports/{export_id}/download")

    out = tmp_path / "export.docx"
    out.write_bytes(download.content)
    doc = Document(str(out))
    assert len(doc.tables) == 0


async def test_list_exports_ordered_newest_first(client, db_session, seeded_draft):
    draft_id = await _analyzed_draft(client, db_session, seeded_draft)
    first = await client.post(f"/api/v1/drafts/{draft_id}/exports", json={"format": "csv"})
    second = await client.post(f"/api/v1/drafts/{draft_id}/exports", json={"format": "json"})

    listing = (await client.get(f"/api/v1/drafts/{draft_id}/exports")).json()
    assert len(listing) == 2
    assert listing[0]["id"] == second.json()["id"]
    assert listing[1]["id"] == first.json()["id"]


async def test_download_unknown_export_404s(client, db_session):
    import uuid

    resp = await client.get(f"/api/v1/exports/{uuid.uuid4()}/download")
    assert resp.status_code == 404


async def test_insertion_mode_direct_produces_no_tracked_changes(
    client, db_session, seeded_draft, tmp_path
):
    draft_id = await _analyzed_draft(client, db_session, seeded_draft)
    resp = await client.post(
        f"/api/v1/drafts/{draft_id}/exports",
        json={"format": "docx", "insertion_mode": "direct"},
    )
    export_id = resp.json()["id"]
    download = await client.get(f"/api/v1/exports/{export_id}/download")

    out = tmp_path / "export.docx"
    out.write_bytes(download.content)
    doc = Document(str(out))
    assert len(doc.element.body.xpath(".//w:ins")) == 0
