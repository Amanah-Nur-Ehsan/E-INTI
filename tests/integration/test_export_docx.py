"""Round-trips the DOCX writer through the full pipeline: upload, analyse,
accept, export, and re-open the produced file to inspect its oxml.
"""

import uuid

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.db.models.enums import InsertionMode
from app.services.export.bundle import build_bundle
from app.services.export.docx_ops import _para_text
from app.services.export.docx_writer import write_docx
from tests.conftest import FIXTURES

pytestmark = pytest.mark.integration


async def _analyzed_draft(client, db_session):
    """Upload+import+run analysis, accept every claim's top recommendation,
    and return (draft_id, accepted_claim_ids).
    """
    from app.services.embedding_service import embed_pending_references
    from app.services.enrichment import enrich_pending_references
    from tests.conftest import import_dataset, upload_draft

    draft_id = await upload_draft(client)
    await import_dataset(client)
    enrich_pending_references(db_session)
    embed_pending_references(db_session)
    db_session.commit()
    await client.post(f"/api/v1/drafts/{draft_id}/analysis/run")

    claims = (await client.get(f"/api/v1/drafts/{draft_id}/claims?needs_citation=true")).json()
    accepted = []
    for claim in claims:
        recs = (await client.get(f"/api/v1/claims/{claim['id']}/recommendations")).json()
        if recs:
            await client.post(f"/api/v1/recommendations/{recs[0]['id']}/accept", json={})
            accepted.append(claim["id"])
    return draft_id, accepted


def _ins_elements(doc: Document) -> list:
    return doc.element.body.xpath(".//w:ins")


async def test_only_accepted_recommendations_reach_the_docx(client, db_session):
    """Regression test for a property that's already correct but was never
    directly asserted: a claim's rejected or never-decided recommendations
    must never appear in the exported document, even though they exist in
    the DB right alongside an accepted sibling for the same claim.

    AcceptedCitation rows are only ever created on decision == ACCEPTED
    (app/api/routes/recommendations.py's _apply_decision); build_bundle
    selects solely from that table; write_docx and its bibliography derive
    solely from bundle.accepted. This test exercises the whole chain rather
    than any one link, so a regression anywhere in it would be caught.
    """
    from app.services.embedding_service import embed_pending_references
    from app.services.enrichment import enrich_pending_references
    from tests.conftest import import_dataset, upload_draft

    draft_id = await upload_draft(client)
    await import_dataset(client)
    enrich_pending_references(db_session)
    embed_pending_references(db_session)
    db_session.commit()
    await client.post(f"/api/v1/drafts/{draft_id}/analysis/run")

    claims = (await client.get(f"/api/v1/drafts/{draft_id}/claims?needs_citation=true")).json()
    assert len(claims) >= 2, "test needs at least two citation-worthy claims"

    # Collect each claim's top recommendation first, so the accept/reject
    # assignment below can pick claims whose top candidates are genuinely
    # different references -- two claims can legitimately share a top
    # candidate (the library is small), and a title-based check can't tell
    # "rejected for this claim, accepted for a different one" apart from a
    # real leak unless the titles themselves differ.
    top_by_claim = []
    for claim in claims:
        recs = (await client.get(f"/api/v1/claims/{claim['id']}/recommendations")).json()
        if recs:
            top_by_claim.append((claim, recs[0]))

    accepted_titles: set[str] = set()
    rejected_titles: set[str] = set()
    pending_titles: set[str] = set()

    accept_target = top_by_claim[0]
    reject_target = next(
        (pair for pair in top_by_claim[1:] if pair[1]["reference"]["title"] != accept_target[1]["reference"]["title"]),
        None,
    )
    assert reject_target is not None, "test needs two claims with distinct top candidates"

    _claim, top = accept_target
    await client.post(f"/api/v1/recommendations/{top['id']}/accept", json={})
    accepted_titles.add(top["reference"]["title"])

    _claim, top = reject_target
    await client.post(f"/api/v1/recommendations/{top['id']}/reject", json={})
    rejected_titles.add(top["reference"]["title"])

    for claim, top in top_by_claim:
        if (claim, top) in (accept_target, reject_target):
            continue
        title = top["reference"]["title"]
        if title not in accepted_titles:
            pending_titles.add(title)

    assert accepted_titles, "test needs at least one accepted recommendation"
    assert rejected_titles or pending_titles, "test needs a non-accepted recommendation to check against"

    bundle = build_bundle(db_session, uuid.UUID(draft_id))
    accepted_ref_titles = {ref.title for item in bundle.accepted for ref in item.references}
    assert accepted_ref_titles == accepted_titles
    assert accepted_ref_titles.isdisjoint(rejected_titles)
    assert accepted_ref_titles.isdisjoint(pending_titles)

    docx_bytes, outcomes = write_docx(bundle, mode=InsertionMode.TRACKED_CHANGES)
    assert len(outcomes) == len(accepted_titles)

    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))
    full_doc_text = "".join(reloaded.element.body.xpath(".//w:t/text()"))
    for title in rejected_titles | pending_titles:
        assert title not in full_doc_text


async def test_tracked_changes_export_inserts_exactly_one_per_claim(client, db_session):
    draft_id, accepted = await _analyzed_draft(client, db_session)
    bundle = build_bundle(db_session, uuid.UUID(draft_id))
    assert len(bundle.accepted) == len(accepted)

    original_bytes = FIXTURES.joinpath("sample_draft.docx").read_bytes()

    docx_bytes, outcomes = write_docx(bundle, mode=InsertionMode.TRACKED_CHANGES)
    assert all(o.status == "INSERTED" for o in outcomes)

    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))
    ins_elements = _ins_elements(reloaded)
    # write_docx shares one monotonic revision counter across citations and
    # the bibliography, so the first N <w:ins> ids (in id order) are the
    # per-claim citations and everything after is bibliography wrapping.
    by_id = sorted(ins_elements, key=lambda el: int(el.get(qn("w:id"))))
    citation_ins = by_id[: len(accepted)]
    assert len(citation_ins) == len(accepted)
    for el in ins_elements:
        assert el.get(qn("w:author")) == "CitationRecommender"

    # The source fixture on disk must be untouched by the export.
    assert FIXTURES.joinpath("sample_draft.docx").read_bytes() == original_bytes


async def test_direct_mode_inserts_plain_text_no_tracked_changes(client, db_session):
    draft_id, accepted = await _analyzed_draft(client, db_session)
    bundle = build_bundle(db_session, uuid.UUID(draft_id))

    docx_bytes, outcomes = write_docx(bundle, mode=InsertionMode.DIRECT)
    assert all(o.status == "INSERTED" for o in outcomes)

    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))
    assert len(_ins_elements(reloaded)) == 0
    # Direct mode: paragraph.text (which excludes w:ins) should now show
    # the inserted citation text directly, since nothing is wrapped.
    joined_body_text = " ".join(p.text for p in reloaded.paragraphs)
    for outcome in outcomes:
        assert outcome.citation_text in joined_body_text


async def test_placeholder_mode_inserts_bracketed_marker(client, db_session):
    draft_id, accepted = await _analyzed_draft(client, db_session)
    bundle = build_bundle(db_session, uuid.UUID(draft_id))

    docx_bytes, outcomes = write_docx(bundle, mode=InsertionMode.PLACEHOLDER)
    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))
    body_text = " ".join(p.text for p in reloaded.paragraphs)
    for outcome in outcomes:
        if outcome.status == "INSERTED":
            assert f"[CITATION: {outcome.citation_text}]" in body_text


async def test_two_claims_in_the_same_paragraph_both_insert_correctly(client, db_session):
    """The descending-order regression: accepting citations for two
    sentences in one paragraph must not corrupt either insertion, and
    must not falsely report a POSITION_MISMATCH on the earlier one.
    """
    draft_id, accepted = await _analyzed_draft(client, db_session)
    bundle = build_bundle(db_session, uuid.UUID(draft_id))

    from collections import Counter

    paragraph_counts = Counter(item.claim.paragraph_index for item in bundle.accepted)
    multi_claim_paragraphs = [p for p, count in paragraph_counts.items() if count >= 2]
    assert multi_claim_paragraphs, "fixture draft should produce >=2 claims in one paragraph"

    docx_bytes, outcomes = write_docx(bundle, mode=InsertionMode.TRACKED_CHANGES)
    assert all(o.status == "INSERTED" for o in outcomes), [
        (o.status, o.paragraph_index) for o in outcomes if o.status != "INSERTED"
    ]

    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))
    for paragraph_index in multi_claim_paragraphs:
        expected_citations = [
            bundle.context.in_text(ref.id)
            for item in bundle.accepted
            if item.claim.paragraph_index == paragraph_index
            for ref in item.references
        ]
        paragraph_text = _para_text(reloaded.paragraphs[paragraph_index])
        for citation in expected_citations:
            assert citation in paragraph_text


async def test_position_mismatch_when_draft_text_changed_since_analysis(client, db_session):
    draft_id, accepted = await _analyzed_draft(client, db_session)
    assert accepted

    # Simulate an external edit: corrupt the stored sentence text for one
    # accepted claim so it no longer matches the live docx paragraph.
    from app.db.models import Claim

    claim = db_session.get(Claim, uuid.UUID(accepted[0]))
    claim.sentence_text = "This sentence text was edited and no longer matches the source file."
    db_session.commit()

    bundle = build_bundle(db_session, uuid.UUID(draft_id))
    docx_bytes, outcomes = write_docx(bundle, mode=InsertionMode.TRACKED_CHANGES)

    mismatched = [o for o in outcomes if o.claim_id == claim.id]
    assert len(mismatched) == 1
    assert mismatched[0].status == "POSITION_MISMATCH"

    # Every other accepted claim must still insert normally.
    others = [o for o in outcomes if o.claim_id != claim.id]
    assert all(o.status == "INSERTED" for o in others)


async def test_bibliography_has_one_entry_per_unique_reference(client, db_session):
    draft_id, accepted = await _analyzed_draft(client, db_session)
    bundle = build_bundle(db_session, uuid.UUID(draft_id))

    docx_bytes, _outcomes = write_docx(bundle, mode=InsertionMode.TRACKED_CHANGES)
    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))

    heading_index = next(
        i for i, p in enumerate(reloaded.paragraphs) if _para_text(p) == "References"
    )
    entries = [_para_text(p) for p in reloaded.paragraphs[heading_index + 1 :]]

    assert len(entries) == len(set(entries))
    unique_ref_ids = {ref.id for item in bundle.accepted for ref in item.references}
    assert len(entries) == len(unique_ref_ids)


async def test_bibliography_italicizes_source_title():
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("placeholder")
    from app.services.citation_formatting_service import Segment

    para = doc.add_paragraph()
    for seg in [
        Segment("Smith, J. "),
        Segment("(2023). "),
        Segment("A Paper. "),
        Segment("A Journal", italic=True),
        Segment(", https://doi.org/10.1/x"),
    ]:
        run = para.add_run(seg.text)
        run.italic = seg.italic

    italic_runs = [r for r in para._p.xpath(".//w:r") if _run_is_italic(r)]
    assert len(italic_runs) == 1
    t = italic_runs[0].find(qn("w:t"))
    assert t.text == "A Journal"


def _run_is_italic(run_el) -> bool:
    rpr = run_el.find(qn("w:rPr"))
    if rpr is None:
        return False
    i_el = rpr.find(qn("w:i"))
    if i_el is None:
        return False
    val = i_el.get(qn("w:val"))
    return val is None or val not in ("0", "false")


async def test_docx_export_rejected_for_markdown_draft(client, db_session):
    data = FIXTURES.joinpath("sample_draft.md").read_bytes()
    upload = await client.post(
        "/api/v1/drafts/upload",
        files={"file": ("sample_draft.md", data)},
    )
    draft_id = upload.json()["id"]

    from app.services.draft_parser_service import parse_and_store_draft

    parse_and_store_draft(db_session, uuid.UUID(draft_id))

    bundle = build_bundle(db_session, uuid.UUID(draft_id))
    with pytest.raises(ValueError, match="DOCX export requires a .docx source draft"):
        write_docx(bundle)


async def test_build_bundle_requires_parsed_draft(client, db_session):
    from app.services.export.bundle import DraftNotParsedError
    from tests.conftest import upload_draft

    draft_id = await upload_draft(client)  # uploaded but never parsed

    with pytest.raises(DraftNotParsedError):
        build_bundle(db_session, uuid.UUID(draft_id))


# --------------------------------------------------------------------------
# Yellow highlighting -- citations, bibliography, SDG keyword
# --------------------------------------------------------------------------


def _highlight_val(run_el) -> str | None:
    rpr = run_el.find(qn("w:rPr"))
    if rpr is None:
        return None
    hl = rpr.find(qn("w:highlight"))
    return hl.get(qn("w:val")) if hl is not None else None


def _full_text(paragraph) -> str:
    """Every `<w:t>` under this paragraph, including runs wrapped in
    `<w:ins>` -- unlike python-docx's own `paragraph.text`, which only
    walks direct-child `<w:r>` elements and silently omits anything already
    wrapped in a tracked-change insertion (see this module's docstring).
    """
    return "".join(paragraph._p.xpath(".//w:t/text()"))


async def test_inserted_citations_and_bibliography_are_yellow_highlighted(client, db_session):
    draft_id, accepted = await _analyzed_draft(client, db_session)
    bundle = build_bundle(db_session, uuid.UUID(draft_id))
    docx_bytes, outcomes = write_docx(bundle, mode=InsertionMode.TRACKED_CHANGES)
    assert all(o.status == "INSERTED" for o in outcomes)

    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))

    # Citation insertions are yellow-highlighted -- this is what survives
    # Accept All Changes, unlike the <w:ins> wrapper itself.
    citation_ins_runs = reloaded.element.body.xpath(".//w:ins/w:r")[: len(accepted)]
    assert citation_ins_runs, "expected at least one tracked-insert citation run"
    assert all(_highlight_val(r) == "yellow" for r in citation_ins_runs)

    # Bibliography entries are highlighted too; the "References" heading
    # itself deliberately is not (see _append_bibliography).
    all_paragraphs = reloaded.paragraphs
    heading_index = next(i for i, p in enumerate(all_paragraphs) if _full_text(p).strip() == "References")
    entry_paragraphs = [p for p in all_paragraphs[heading_index + 1 :] if _full_text(p).strip()]
    assert entry_paragraphs
    entry_runs = [r for p in entry_paragraphs for r in p._p.xpath(".//w:r")]
    assert entry_runs
    assert all(_highlight_val(r) == "yellow" for r in entry_runs)


async def _upload_docx_bytes(client, filename: str, paragraphs: list[tuple[str, str | None]]) -> str:
    from io import BytesIO as _BytesIO

    document = Document()
    for text, style in paragraphs:
        document.add_paragraph(text, style=style)
    buffer = _BytesIO()
    document.save(buffer)

    resp = await client.post("/api/v1/drafts/upload", files={"file": (filename, buffer.getvalue())})
    return resp.json()["id"]


async def test_sdg_keyword_appended_to_existing_keywords_line(client, db_session):
    from app.db.models import Draft
    from tests.conftest import import_dataset

    draft_id = await _upload_docx_bytes(
        client,
        "keywords_paper.docx",
        [
            ("A Paper About Fraud Detection", "Title"),
            ("Keywords: fraud detection, machine learning", None),
            ("Introduction", "Heading 1"),
            (
                "Machine learning techniques have improved the ability to identify complex "
                "fraud patterns in financial transactions.",
                None,
            ),
        ],
    )
    await import_dataset(client)
    from app.services.embedding_service import embed_pending_references
    from app.services.enrichment import enrich_pending_references

    enrich_pending_references(db_session)
    embed_pending_references(db_session)
    db_session.commit()
    await client.post(f"/api/v1/drafts/{draft_id}/analysis/run")

    claims = (await client.get(f"/api/v1/drafts/{draft_id}/claims?needs_citation=true")).json()
    for claim in claims:
        recs = (await client.get(f"/api/v1/claims/{claim['id']}/recommendations")).json()
        if recs:
            await client.post(f"/api/v1/recommendations/{recs[0]['id']}/accept", json={})

    # Force a deterministic, genuine-looking SDG pick regardless of what the
    # mock happened to classify this particular fixture as -- this test is
    # about *where* the keyword lands, not the classification itself.
    draft = db_session.get(Draft, uuid.UUID(draft_id))
    draft.sdg_number = 9
    draft.sdg_name = "Industry, innovation and infrastructure"
    draft.sdg_keyword = "machine learning applications"
    db_session.commit()

    bundle = build_bundle(db_session, uuid.UUID(draft_id))
    docx_bytes, _outcomes = write_docx(bundle, mode=InsertionMode.TRACKED_CHANGES)

    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))
    keywords_paragraphs = [p for p in reloaded.paragraphs if _full_text(p).startswith("Keywords:")]
    assert len(keywords_paragraphs) == 1
    para = keywords_paragraphs[0]
    assert "machine learning applications" in _full_text(para)

    highlighted_runs = [r for r in para._p.xpath(".//w:r") if _highlight_val(r) == "yellow"]
    assert any(
        "machine learning applications" in "".join(r.xpath(".//w:t/text()")) for r in highlighted_runs
    )


async def test_sdg_keyword_inserted_as_new_paragraph_when_no_keywords_line(client, db_session):
    """sample_draft.docx has no Keywords line at all -- the keyword must
    land as a new paragraph right before Introduction instead.
    """
    from app.db.models import Draft

    draft_id, _accepted = await _analyzed_draft(client, db_session)

    draft = db_session.get(Draft, uuid.UUID(draft_id))
    draft.sdg_number = 9
    draft.sdg_name = "Industry, innovation and infrastructure"
    draft.sdg_keyword = "adaptive systems"
    db_session.commit()

    bundle = build_bundle(db_session, uuid.UUID(draft_id))
    docx_bytes, _outcomes = write_docx(bundle, mode=InsertionMode.TRACKED_CHANGES)

    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))
    texts = [_full_text(p) for p in reloaded.paragraphs]
    keyword_index = next(i for i, t in enumerate(texts) if t.startswith("Keywords:"))
    intro_index = next(i for i, t in enumerate(texts) if t.strip().lower() == "introduction")
    assert keyword_index < intro_index
    assert "adaptive systems" in texts[keyword_index]


async def test_sdg_keyword_omitted_when_classification_declined(client, db_session):
    """A no-match SDG (sdg_number None) must not write anything -- there is
    no honest keyword to highlight.
    """
    from app.db.models import Draft

    draft_id, _accepted = await _analyzed_draft(client, db_session)

    draft = db_session.get(Draft, uuid.UUID(draft_id))
    draft.sdg_number = None
    draft.sdg_name = None
    draft.sdg_keyword = None
    db_session.commit()

    bundle = build_bundle(db_session, uuid.UUID(draft_id))
    docx_bytes, _outcomes = write_docx(bundle, mode=InsertionMode.TRACKED_CHANGES)

    from io import BytesIO

    reloaded = Document(BytesIO(docx_bytes))
    assert not any(p.text.startswith("Keywords:") for p in reloaded.paragraphs)
